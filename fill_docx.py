import docx

def fill_docx(filename, output_filename):
    doc = docx.Document(filename)
    
    visao_geral = """O Dupla Cor é uma aplicação Web completa com objetivo de otimizar o gerenciamento de uma boutique de esmaltes, garantindo controle inteligente de inventário e vendas no e-commerce.

Escopo do Sistema:
- E-Commerce: Catálogo com busca, filtros, vitrine dinâmica e carrinho persistente.
- Controle FEFO: Rastreabilidade total e controle de validade pelo algoritmo First Expired, First Out.
- Painel Administrativo: Gestão de catálogo, entrada de Lotes, auditoria e pedidos.

Exclusão de Escopo:
- Integração direta com hardware físico de frente de caixa.
- Gestão de RH.
- Emissão de notas fiscais."""

    arquitetura = """O sistema adota uma arquitetura MVC (Model-View-Controller) para o Backend (em Java) combinada com Single Page Application (SPA) no Frontend.
Pacotes:
- Model: Classes POO (Produto, Lote, Pedido).
- DAO: Persistência JDBC.
- Controller: Regras FEFO e API REST.
- View: SPA em HTML5, CSS3, JS.

DuplaCor/
├── Dockerfile & docker-compose.yml
├── database/schema.sql
├── web/ (Frontend Web SPA: HTML, CSS, JS)
└── src/ (Backend Java)
    ├── model/ (Entidades)
    ├── dao/ (JDBC)
    ├── controller/ (Regras de negócio e API REST)
    └── server/ (Servidor HTTP nativo)"""

    modelo_bd = """Tecnologia: MySQL 8.0
Hospedagem: Docker (Local)
Responsabilidades: Persistência dos dados, Integridade referencial, Consultas SQL (Transações seguras)."""

    design = """Componentes:
- Átomos: Button (.btn-primary), Input, BadgeStatus (badge-fefo)
- Moléculas: FormField, ProductCard
- Organismos: Header/Nav, ProductGrid
- Páginas: HomePage, AdminDashboard

Paleta de Cores:
Primária: #9b72cf
Primária Dark: #5c3d99
Texto Principal: #2b1e3a
Background Champagne: #faf7fc

Tipografia: Plus Jakarta Sans / Cormorant Garamond

Responsividade: Desktop e Mobile adaptativos."""

    fluxo = """1. Usuário acessa o sistema (a SPA index.html carrega a #home).
2. O Frontend dispara GET /api/produtos no servidor HTTP nativo.
3. Enquanto aguarda: tela permanece com containers pré-renderizados.
4. Após retorno: O estado local é atualizado e os esmaltes com lotes válidos são renderizados dinamicamente."""

    software = """- Tela da Loja (Home e Vitrine): Exibe o banner de introdução e catálogo dinâmico de produtos com lotes válidos.
- Tela do Carrinho com FEFO: Cliente revisa as quantidades e o sistema exibe de quais lotes as unidades serão retiradas.
- Tela do Painel Administrativo: Indicadores do e-commerce (lotes ativos, vencidos) e gestão logística."""

    link = "Repositório: https://github.com/diule/DuplaCor-main"

    for para in doc.paragraphs:
        text = para.text.strip()
        if "Criar texto explicando, escopo" in text:
            para.text = visao_geral
        elif "Criar texto explicando os pacotes do projeto" in text:
            para.text = arquitetura
        elif "Em seguida, inserir a imagem da árvore" in text:
            para.text = ""
        elif "Tecnologia:" in text:
            para.text = modelo_bd
        elif "Supabase" in text:
            para.text = "Docker"
        elif "Dos componentes:" in text:
            para.text = design
        elif "Fluxo de carregamento:" in text:
            para.text = fluxo
        elif "<Inserir telas do Sistema" in text:
            para.text = software
        elif "<Link do projeto no GitHub>" in text:
            para.text = link
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if "Tecnologia:" in text:
                    cell.text = modelo_bd
                elif "Supabase" in text:
                    cell.text = "Docker"
                elif "Dos componentes:" in text:
                    cell.text = design

    doc.save(output_filename)

if __name__ == "__main__":
    fill_docx("Modelo_projeto Final_2026 (2).docx", "Modelo_projeto_Final_Preenchido.docx")
